import io, json
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt
import streamlit as st

# ------------------------------
# ---- VERSION TRACKING ----
# ------------------------------
DATA_DIR = Path(".data")
DATA_DIR.mkdir(exist_ok=True)
VERSION_STORE = DATA_DIR / "version_store.json"
if VERSION_STORE.exists():
    version_tracker = json.loads(VERSION_STORE.read_text())
else:
    version_tracker = {}

# ------------------------------
# ---- GENERAL SERVICES & SCOPES ----
# ------------------------------
GENERAL_SERVICES_RULES = [
    ("default", "Provide preventive maintenance and inspection labor, materials (filters, oil, grease, rags where applicable), and related expenses."),
    ("default", "Perform maintenance in accordance with manufacturer recommendations for covered equipment."),
    ("default", "Replace or clean filters when and where applicable (quarterly)."),
    ("coil", "Clean condenser and evaporator coils annually, or as manufacturer guidelines require."),
    ("refrigerant", "Monitor compressor operating pressures quarterly (if accessible)."),
    ("default", "Inspect belt tension and wear quarterly where belt-driven assemblies are present."),
    ("default", "Test electrical components and connections quarterly."),
    ("default", "Provide seasonal changeover services (spring and/or fall) where applicable."),
    ("boiler", "Monitor boilers for proper temperatures and pressures."),
    ("boiler", "Check boilers for leaks and sediment buildup."),
    ("hydronic", "Inspect pumps listed in the equipment inventory, or directly associated with covered assets, for leaks, vibration, and audible wear."),
    ("refrigerant", "Provide refrigerant conservation services, including certified leak detection, recovery, and recycling to comply with the Clean Air Act and all applicable state/local regulations."),
]

COIL_EQUIP = {"Air Handler","RTUs","CRAC ID","CRAC OD","Fan Coil","Mini-Splits","Water-Cooled Chillers","Condenser","Cooling Towers"}
REFRIG_EQUIP = {"Air Handler","RTUs","CRAC ID","CRAC OD","Fan Coil","Mini-Splits","Water-Cooled Chillers","Condenser"}
BOILER_EQUIP = {"Boilers"}
HYDRONIC_EQUIP = {"Pumps","Boilers","Water-Cooled Chillers","Cooling Towers","Fan Coil","Air Handler"}

SCOPES = {
    "Boilers": ("Boilers:",
        "• Annual: Inspect burner assembly; test flame detection, fuel cut-off, low-water cutoff, and safety valves; inspect fuel system; record operating pressures, gas settings, boiler input/output; calibrate controls; run combustion analysis.",
        "• Quarterly: Test flame detection, operating control, flow switch, and fuel system; inspect pump operation; blow down low-water cutoff; verify flame condition."),
    "Cooling Towers": ("Cooling Towers:",
        "• Annual: Remove debris, flush basin, inspect/clean strainers, lubricate bearings, change gear reducer oil (if equipped), inspect belts/pulleys, check electrical connections and controls, record motor amp draw, verify float valve operation, inspect spray nozzles, check fill media.",
        "• Quarterly: Lubricate bearings, inspect belts/pulleys (if equipped), check controls, record amp draw, verify float valve, inspect/clean spray nozzles."),
    "Pumps": ("Pumps:",
        "• Annual: Verify pump identification, inspect coupling, check for vibration/noise, record motor RLA and actual amps, inspect starter and connections, tighten electrical terminals, lubricate bearings (if applicable), check motor alignment, record voltage, return to service.",
        "• Quarterly: Verify pump identification, inspect for vibration/noise, check coupling, record motor RLA and amps, return to service."),
    "Air Handler": ("Air Handlers:",
        "• Annual: Inspect and clean coils; replace or clean filters (if applicable); lubricate bearings; inspect blower assembly for wear; check fan wheel alignment; inspect and tighten electrical connections; verify damper operation; calibrate control sensors.",
        "• Quarterly: Inspect filters and clean/replace as needed; check coil condition; inspect and adjust belt tension (if applicable); verify motor amperage and operating condition; inspect condensate pans and drains for proper operation."),
    "RTUs": ("Roof Top Units (RTUs):",
        "• Annual: Inspect and clean condenser and evaporator coils; replace or clean air filters (where applicable); inspect and lubricate fan and motor bearings (if applicable); inspect blower assembly for alignment and wear; inspect belts and pulleys for wear and proper tension (if equipped); inspect heat exchangers for cracks or corrosion; check refrigerant charge, pressures, and superheat/subcooling; inspect condensate pans and drains, clean and flush as needed; inspect electrical connections, contactors, relays, and capacitors; test safety controls and high/low pressure switches; verify economizer operation and damper function (if installed); record supply/return air temperatures, amperages, and voltages; calibrate thermostats and sensors as needed; perform combustion analysis on gas-fired sections (if applicable).",
        "• Quarterly: Inspect and clean or replace filters as required; inspect coil condition and clean as needed; inspect belts/pulleys and adjust or replace as necessary; check fan and blower motor operation, record amperage draws; inspect and tighten electrical connections; inspect condensate drains for blockage and proper operation; record refrigerant pressures and system temperatures (if accessible); verify overall unit operation within manufacturer specifications; inspect cabinet condition, economizer (if applicable), and access panels."),
    # (Other equipment types omitted for brevity; you’d paste the full list we finalized earlier here)
}

ALIASES = {k.lower(): k for k in SCOPES.keys()}
ALIASES.update({"ahu":"Air Handler","rtu":"RTUs","crac id":"CRAC ID","crac od":"CRAC OD"})

def canonical_scope_name(raw): return ALIASES.get((raw or "").strip().lower())

def visits_text_from_frequency(freq):
    f = (freq or "").lower()
    if f.startswith("annual"): return "one (1) annual service per year"
    if f.startswith("semi"): return "one (1) operating inspection and one (1) annual service per year"
    if f.startswith("quarter"): return "three (3) operating inspections and one (1) annual service per year"
    return "three (3) operating inspections and one (1) annual service per year"

def billing_text_from_frequency(freq):
    f = (freq or "").lower()
    if f.startswith("annual"): return "Billing will occur annually following completion of the scheduled maintenance visit."
    if f.startswith("semi"): return "Billing will occur semi-annually following completion of each scheduled maintenance visit."
    if f.startswith("quarter"): return "Billing will occur quarterly following completion of each scheduled maintenance visit."
    return "Billing will occur quarterly following completion of each scheduled maintenance visit."

def fraction_from_frequency(freq):
    f = (freq or "").lower()
    if f.startswith("annual"): return "the full annual agreement total"
    if f.startswith("semi"): return "one-half (1/2) of the annual agreement total"
    if f.startswith("quarter"): return "one-fourth (1/4) of the annual agreement total"
    return "one-fourth (1/4) of the annual agreement total"

def parse_intake(file: bytes):
    df = pd.read_excel(io.BytesIO(file), sheet_name="Service Intake", header=None)
    def find_row(label):
        for i in range(300):
            row_vals = [str(x).strip() if pd.notna(x) else "" for x in df.iloc[i, :8]]
            if any(v == label for v in row_vals): return i
        return None
    def get_dates_dict(header_row):
        headers = [str(h).strip() if pd.notna(h) else "" for h in df.iloc[header_row, 0:3]]
        values = df.iloc[header_row+1, 0:3].tolist()
        out={}
        for h,v in zip(headers,values):
            if not h: continue
            if pd.isna(v): out[h]=""
            else:
                try: out[h]=pd.to_datetime(v).date().isoformat()
                except: out[h]=str(v).strip()
        return out
    def get_equipment_rows(header_row):
        headers = [str(h).strip() if pd.notna(h) else "" for h in df.iloc[header_row, 0:6]]
        rows=[]; r=header_row+1
        while r < len(df):
            row_vals = df.iloc[r, 0:6]
            if row_vals.isna().all(): r+=1; continue
            vals=[str(v).strip() if pd.notna(v) else "" for v in row_vals]
            if any(vals): rows.append(dict(zip(headers, vals)))
            r+=1
        return rows
    agree_row = find_row("Preferred Start Date")
    equip_row = find_row("Equipment Type")
    return {"agreement": get_dates_dict(agree_row) if agree_row else {}, "equipment_rows": get_equipment_rows(equip_row) if equip_row else []}

def generate_filename(property_name, start_date, end_date):
    try: start_year = str(pd.to_datetime(start_date).year)
    except: start_year = "XXXX"
    try: end_year = str(pd.to_datetime(end_date).year)
    except: end_year = "XXXX"
    base = f"{property_name} - PM Agreement - {start_year}-{end_year}"
    current = version_tracker.get(base, 0)
    version_tracker[base] = current + 1
    ver_str = f"V1.{current}"
    VERSION_STORE.write_text(json.dumps(version_tracker))
    return f"{base} - {ver_str}.docx"

def build_doc(payload, property_name):
    freq = payload["agreement"].get("Service Frequency","")
    visits_text = visits_text_from_frequency(freq)
    billing_text = billing_text_from_frequency(freq)
    fraction_text = fraction_from_frequency(freq)
    start_date = payload["agreement"].get("Preferred Start Date","[Start Date]")
    end_date = payload["agreement"].get("Preferred End Date","[End Date]")

    present = set()
    for r in payload["equipment_rows"]:
        canon = canonical_scope_name(r.get("Equipment Type",""))
        if canon in SCOPES: present.add(canon)

    filtered=[]
    for tag,text in GENERAL_SERVICES_RULES:
        if tag=="default": filtered.append(text)
        elif tag=="coil" and present & COIL_EQUIP: filtered.append(text)
        elif tag=="refrigerant" and present & REFRIG_EQUIP: filtered.append(text)
        elif tag=="boiler" and present & BOILER_EQUIP: filtered.append(text)
        elif tag=="hydronic" and present & HYDRONIC_EQUIP: filtered.append(text)
    general = [f"{i}.\t{t}" for i,t in enumerate(filtered, start=1)]

    doc = Document()
    style = doc.styles["Normal"]; style.font.name="Calibri"; style.font.size=Pt(11)

    intro = (f"This Preventive Maintenance Program includes {visits_text}. "
             "All maintenance and inspection work will be conducted during regular business hours (M–F, 8:00 AM to 5:00 PM). "
             "A written report of findings, corrective actions, and recommendations will follow each visit. "
             "Corrective actions not covered under this agreement will be quoted for approval prior to performance.")
    doc.add_paragraph(intro); doc.add_paragraph()
    doc.add_paragraph("General Services")
    for line in general: doc.add_paragraph(line)

    if present:
        doc.add_paragraph(); doc.add_paragraph("Equipment-Specific Services")
        for name in sorted(present):
            hdr, ann, qtr = SCOPES[name]
            doc.add_paragraph(hdr); doc.add_paragraph(ann); doc.add_paragraph(qtr)

    doc.add_paragraph()
    doc.add_paragraph("Preferred Client Status:\n\n"
                      "Upon acceptance of this agreement, the customer will be recognized as a Preferred Client. Benefits include:\n"
                      "•\tDiscounted service labor rates: $145.00 per hour during regular business hours (M–F, 8:00 AM – 5:00 PM) and $217.50 per hour for after-hours, weekends, or holidays.\n"
                      "•\tPriority scheduling for emergency service requests.\n"
                      "•\tReduced rates for repair parts, components, regulated material recovery, and disposal services.\n"
                      "•\tWritten service tickets provided after each inspection or repair.")
    doc.add_paragraph()
    doc.add_paragraph(f"Agreement Term:\n\nThis Service Agreement will commence on {start_date} and continue through {end_date}, unless terminated or renewed in accordance with contract terms.")
    doc.add_paragraph()
    doc.add_paragraph("Billing & Payment:\n\n"
                      f"{billing_text}\n\n"
                      f"Each invoice will reflect {fraction_text}, plus applicable taxes.\n\n"
                      "Additional services, repairs, or emergency calls outside the scope of this agreement will be billed separately at Preferred Client rates.")

    filename = generate_filename(property_name, start_date, end_date)
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return filename, bio

# ------------------------------
# ---- STREAMLIT UI ----
# ------------------------------
st.title("PM Agreement Scope Generator")
st.caption("Upload your completed Service Intake Excel to generate a formatted DOCX scope.")

uploaded = st.file_uploader("Upload Service Intake (.xlsx)", type=["xlsx"])
prop_name = st.text_input("Property Name (for filename)", value="")
alpha = st.checkbox("Alphabetize equipment sections", value=True)

if st.button("Generate Scope") and uploaded is not None:
    payload = parse_intake(uploaded.getvalue())
    filename, bio = build_doc(payload, property_name=(prop_name or "Property"))
    st.success("Scope generated.")
    st.download_button("Download DOCX", data=bio, file_name=filename,
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
